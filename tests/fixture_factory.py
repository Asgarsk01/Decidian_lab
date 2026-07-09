from __future__ import annotations

from pathlib import Path


def create_test_image(path: Path) -> Path:
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (900, 420), "white")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((60, 70, 330, 180), radius=20, fill="#dbeafe", outline="#2563eb", width=4)
    draw.rounded_rectangle((570, 70, 840, 180), radius=20, fill="#dcfce7", outline="#16a34a", width=4)
    draw.line((330, 125, 570, 125), fill="#111827", width=6)
    draw.polygon([(550, 110), (575, 125), (550, 140)], fill="#111827")
    draw.text((115, 110), "API Service", fill="#111827")
    draw.text((635, 110), "Order Queue", fill="#111827")
    draw.text((180, 270), "All payment retries must be capped at three attempts.", fill="#111827")
    image.save(path, "PNG")
    return path


def create_text_pdf(path: Path, image_path: Path) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    styles = getSampleStyleSheet()
    story = [
        Paragraph("RFC-017: Reliable Payment Processing", styles["Title"]),
        Paragraph("1. Status", styles["Heading1"]),
        Paragraph("Approved by the payments team on 2026-06-01.", styles["BodyText"]),
        Paragraph("2. Decision", styles["Heading1"]),
        Paragraph(
            "Payment requests shall be placed on the Order Queue and processed asynchronously. "
            "Retries must be capped at three attempts.",
            styles["BodyText"],
        ),
        Spacer(1, 12),
        ListFlowable(
            [
                ListItem(Paragraph("Use idempotency keys for every request.", styles["BodyText"])),
                ListItem(Paragraph("Reject duplicate completed payments.", styles["BodyText"])),
            ],
            bulletType="bullet",
        ),
        Paragraph("3. Retry policy", styles["Heading1"]),
    ]
    table = Table(
        [
            ["Failure", "Maximum retries", "Final action"],
            ["Network timeout", "3", "Move to dead-letter queue"],
            ["Validation error", "0", "Reject immediately"],
        ]
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightblue),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.extend(
        [
            table,
            Spacer(1, 20),
            Paragraph("4. Architecture", styles["Heading1"]),
            Image(str(image_path), width=6.4 * inch, height=3.0 * inch),
        ]
    )
    SimpleDocTemplate(str(path), pagesize=A4).build(story)
    return path


def create_docx(path: Path, image_path: Path) -> Path:
    from docx import Document
    from docx.shared import Inches

    document = Document()
    document.add_heading("SRS: Order Processing", level=0)
    document.add_heading("1. Approved requirements", level=1)
    document.add_paragraph(
        "The system shall preserve the idempotency key for every order."
    )
    document.add_paragraph(
        "The service shall never process a completed payment twice.",
        style="List Bullet",
    )
    document.add_heading("2. Retry constraints", level=1)
    table = document.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Failure"
    table.rows[0].cells[1].text = "Retries"
    table.rows[0].cells[2].text = "Action"
    row = table.add_row().cells
    row[0].text = "Timeout"
    row[1].text = "3"
    row[2].text = "Dead-letter queue"
    document.add_heading("3. Architecture", level=1)
    document.add_picture(str(image_path), width=Inches(6))
    document.save(path)
    return path


def create_scanned_pdf(path: Path) -> Path:
    from PIL import Image, ImageDraw

    page = Image.new("RGB", (1654, 2339), "white")
    draw = ImageDraw.Draw(page)
    draw.text((120, 140), "APPROVED SECURITY REQUIREMENT", fill="black")
    draw.text((120, 260), "All administrator sessions must require MFA.", fill="black")
    draw.text((120, 340), "MFA failures must never allow access.", fill="black")
    draw.rectangle((110, 460, 1500, 850), outline="black", width=4)
    draw.text((150, 510), "Requirement | Owner | Criticality", fill="black")
    draw.text((150, 610), "MFA required | Security TL | Critical", fill="black")
    page.save(path, "PDF", resolution=150.0)
    return path


def create_all_fixtures(directory: Path) -> dict[str, Path]:
    directory.mkdir(parents=True, exist_ok=True)
    image = create_test_image(directory / "architecture.png")
    return {
        "pdf": create_text_pdf(directory / "sample-rfc.pdf", image),
        "docx": create_docx(directory / "sample-srs.docx", image),
        "scanned": create_scanned_pdf(directory / "sample-scanned.pdf"),
    }

