from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from decidian_docling.canonical import build_canonical_document, canonical_markdown
from decidian_docling.clean_chunking import build_clean_chunks, build_review_queue
from decidian_docling.config import GeminiSettings, get_gemini_settings
from decidian_docling.gemini_review import apply_verified_overlays, run_gemini_review


class WordTokenizer:
    def count_tokens(self, text: str) -> int:
        return len(text.split())


def _number(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    p_pr.append(num_pr)


def test_docx_canonicalization_protects_code_and_demotes_records(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("10. Deployment", level=1)
    document.add_heading("10.2. Docker Compose Configuration", level=2)
    document.add_paragraph("# docker-compose.yml - Service Overview")
    document.add_paragraph("services:")
    document.add_paragraph("# Application Service")
    document.add_paragraph("app:")
    record = document.add_paragraph("Email must be valid email format", style="Heading 3")
    _number(record)
    document.add_heading("7.6.2. 8D Report Structure", level=4)
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Permission"
    table.rows[0].cells[1].text = "Description"
    table.rows[1].cells[0].text = "modules:read"
    table.rows[1].cells[1].text = "Read modules"
    document.save(source)

    result = build_canonical_document(source, {"texts": [], "pictures": []}, {}, {}, [])
    blocks = result.document["blocks"]
    code = next(block for block in blocks if block["type"] == "code")
    record_block = next(block for block in blocks if "Email must" in block.get("text", ""))
    report_heading = next(block for block in blocks if "8D Report" in block.get("text", ""))

    assert "# Application Service" in code["text"]
    assert record_block["type"] == "list_item"
    assert report_heading["level"] == 3
    assert "```yaml" in canonical_markdown(result.document)
    table_block = next(block for block in blocks if block["type"] == "table")
    assert any(ref.startswith("docx://paragraphs/") for ref in code["source_refs"])
    assert table_block["source_refs"] == ["docx://tables/0"]


def test_clean_chunks_keep_table_rows_and_field_values_atomic() -> None:
    canonical = {
        "blocks": [
            {
                "id": "block-1",
                "type": "paragraph",
                "text": "Permission =",
                "section_path": ["Authorization"],
                "source_refs": ["#/texts/1"],
                "page_numbers": [],
                "integrity_status": "verified",
                "integrity_finding_ids": [],
            },
            {
                "id": "block-2",
                "type": "paragraph",
                "text": "modules:read",
                "section_path": ["Authorization"],
                "source_refs": ["#/texts/2"],
                "page_numbers": [],
                "integrity_status": "verified",
                "integrity_finding_ids": [],
            },
            {
                "id": "block-3",
                "type": "table",
                "headers": ["Field", "Data Type"],
                "rows": [["id", "UUID"], ["tenant_id", "UUID"]],
                "section_path": ["Database"],
                "source_refs": ["#/tables/1"],
                "page_numbers": [3],
                "integrity_status": "verified",
                "integrity_finding_ids": [],
            },
        ]
    }

    chunks, config = build_clean_chunks(canonical, tokenizer=WordTokenizer())

    assert "Permission = modules:read" in chunks[0]["text"]
    assert [chunk["content_type"] for chunk in chunks[1:]] == ["table_row", "table_row"]
    assert "Field = id" in chunks[1]["text"] and "Data Type = UUID" in chunks[1]["text"]
    assert config["excluded_blocks"] == 0


def test_pdf_integrity_finding_becomes_targeted_candidate(tmp_path: Path) -> None:
    source = tmp_path / "sample.pdf"
    source.write_bytes(b"%PDF-1.4\n")
    document_data = {
        "body": {"children": [{"$ref": "#/tables/0"}]},
        "texts": [],
        "groups": [],
        "pictures": [],
        "tables": [{
            "self_ref": "#/tables/0",
            "prov": [{"page_no": 2, "bbox": {"l": 1, "t": 9, "r": 8, "b": 2}}],
            "data": {
                "num_rows": 2,
                "num_cols": 2,
                "table_cells": [
                    {"start_row_offset_idx": 0, "start_col_offset_idx": 0, "text": "1"},
                    {"start_row_offset_idx": 0, "start_col_offset_idx": 1, "text": "2"},
                    {"start_row_offset_idx": 1, "start_col_offset_idx": 0, "text": "id"},
                    {"start_row_offset_idx": 1, "start_col_offset_idx": 1, "text": "UUID"},
                ],
            },
        }],
    }
    semantic = {
        "findings": [{
            "id": "si-0001",
            "category": "ambiguous_numeric_table_header",
            "status": "review_required",
            "blocks_llm_readiness": True,
            "source_table_refs": ["#/tables/0"],
            "source_refs": [],
            "pages": [2],
        }]
    }

    result = build_canonical_document(source, document_data, semantic, {}, [])

    table = result.document["blocks"][0]
    assert table["integrity_status"] == "review_required"
    assert result.candidates[0]["kind"] == "structural_ambiguity"
    assert result.candidates[0]["page_numbers"] == [2]


def test_missing_gemini_key_queues_candidate_without_failing(tmp_path: Path) -> None:
    settings = GeminiSettings(
        enabled=True,
        api_key=None,
        model="gemini-test",
        timeout_seconds=30,
        max_retries=0,
        max_concurrency=1,
    )
    candidate = {
        "id": "ai-0001",
        "kind": "diagram",
        "block_id": "block-1",
        "source_refs": ["#/pictures/0"],
        "page_numbers": [],
        "section_path": ["Architecture"],
        "ambiguity_reasons": ["content_bearing_picture"],
    }

    report, results, supplements = run_gemini_review([candidate], tmp_path, settings, tmp_path / "cache")
    queue = build_review_queue(
        {"blocks": [{"id": "block-1", "section_path": ["Architecture"]}]},
        [candidate],
        results,
    )

    assert report["status"] == "not_configured"
    assert supplements == []
    assert queue[0]["status"] == "not_reviewed"


def test_only_fully_verified_structural_overlay_is_applied() -> None:
    canonical = {
        "blocks": [
            {
                "id": "block-1",
                "type": "paragraph",
                "level": None,
                "integrity_status": "ambiguous",
            }
        ]
    }
    candidates = [{"id": "ai-0001", "kind": "structural_ambiguity", "block_id": "block-1"}]
    apply_verified_overlays(
        canonical,
        candidates,
        {
            "ai-0001": {
                "status": "verified",
                "extraction": {"corrected_block_type": "heading", "corrected_level": 2},
            }
        },
    )

    assert canonical["blocks"][0]["type"] == "heading"
    assert canonical["blocks"][0]["level"] == 2
    assert canonical["blocks"][0]["integrity_status"] == "ai_verified"


def test_two_pass_gemini_promotes_only_verified_claims(monkeypatch, tmp_path: Path) -> None:
    import json
    from google import genai

    class Response:
        usage = None

        def __init__(self, output_text: str) -> None:
            self.output_text = output_text

    requests = []

    class Interactions:
        def create(self, **kwargs):
            requests.append(kwargs)
            prompt = kwargs["input"][0]["text"]
            if "Independently verify" in prompt:
                return Response(json.dumps({
                    "candidate_id": "ai-0001",
                    "verdicts": [{
                        "claim_id": "c-0001",
                        "verdict": "verified",
                        "evidence": [{"label": "API to database", "box_2d": [10, 10, 50, 50]}],
                        "corrected_statement": None,
                        "rationale": "The directed connector and labels are visible.",
                    }],
                    "conflicts": [],
                    "unresolved": [],
                }))
            return Response(json.dumps({
                "candidate_id": "ai-0001",
                "classification": "architecture_diagram",
                "visible_labels": [{"label": "API", "box_2d": [10, 10, 20, 20]}],
                "components": ["API", "Database"],
                "relationships": ["API -> Database"],
                "corrected_block_type": None,
                "corrected_level": None,
                "claims": [{
                    "claim_id": "c-0001",
                    "claim_type": "relationship",
                    "statement": "The API connects to the database.",
                    "evidence": [{"label": "API to database", "box_2d": [10, 10, 50, 50]}],
                }],
                "unresolved": [],
            }))

    class Client:
        interactions = Interactions()

    monkeypatch.setattr(genai, "Client", lambda **kwargs: Client())
    settings = GeminiSettings(True, "secret", "gemini-test", 30, 0, 1)
    candidate = {
        "id": "ai-0001",
        "kind": "diagram",
        "block_id": "block-1",
        "source_refs": ["#/pictures/0"],
        "page_numbers": [1],
        "section_path": ["Architecture"],
        "ambiguity_reasons": ["content_bearing_picture"],
    }

    report, results, supplements = run_gemini_review([candidate], tmp_path, settings, tmp_path / "cache")

    assert report["status"] == "success"
    assert results["ai-0001"]["status"] == "verified"
    assert supplements[0]["statement"] == "The API connects to the database."
    assert all(request["api_version"] == "v1beta" for request in requests)
    assert results["ai-0001"]["timings"]["total_seconds"] >= 0
    assert "secret" not in json.dumps(report)


def test_systemic_gemini_failure_opens_circuit_and_preserves_evidence(
    monkeypatch,
    tmp_path: Path,
) -> None:
    from google import genai

    pictures = tmp_path / "pictures"
    pictures.mkdir()
    for index in range(1, 4):
        (pictures / f"picture-{index:04d}.png").write_bytes(
            f"test-image-{index}".encode()
        )
    calls = []

    class Interactions:
        def create(self, **kwargs):
            calls.append(kwargs)
            raise RuntimeError("Error code: 404 - model not found for key secret")

    class Client:
        interactions = Interactions()

    monkeypatch.setattr(genai, "Client", lambda **kwargs: Client())
    settings = GeminiSettings(True, "secret", "gemini-test", 30, 0, 2)
    candidates = [
        {
            "id": f"ai-{index:04d}",
            "kind": "diagram",
            "block_id": f"block-{index}",
            "source_refs": [f"#/pictures/{index - 1}"],
            "page_numbers": [],
            "section_path": ["Architecture"],
            "picture_file": f"picture-{index:04d}.png",
            "ambiguity_reasons": ["content_bearing_picture"],
        }
        for index in range(1, 4)
    ]

    report, results, supplements = run_gemini_review(
        candidates,
        tmp_path,
        settings,
        tmp_path / "cache",
    )
    queue = build_review_queue(
        {
            "blocks": [
                {"id": f"block-{index}", "section_path": ["Architecture"]}
                for index in range(1, 4)
            ]
        },
        candidates,
        results,
    )

    assert report["status"] == "failed"
    assert report["circuit_breaker_triggered"] is True
    assert len(calls) == 1
    assert results["ai-0001"]["attempted"] is True
    assert results["ai-0001"]["evidence_sha256"]
    assert results["ai-0002"]["attempted"] is False
    assert results["ai-0002"]["blocked_by_candidate_id"] == "ai-0001"
    assert candidates[0]["evidence_path"].startswith("ai_evidence/")
    assert queue[0]["evidence_path"] == candidates[0]["evidence_path"]
    assert supplements == []
    assert "secret" not in json.dumps(report)


def test_guarded_review_caps_candidates_skips_empty_claim_verification_and_emits_events(
    monkeypatch,
    tmp_path: Path,
) -> None:
    from google import genai

    class Usage:
        def model_dump(self, **kwargs):
            return {
                "total_input_tokens": 1_000,
                "total_output_tokens": 200,
                "total_thought_tokens": 300,
                "total_tokens": 1_500,
            }

    requests = []

    class Response:
        usage = Usage()

        def __init__(self, candidate_id: str) -> None:
            self.output_text = json.dumps({
                "candidate_id": candidate_id,
                "classification": "unreadable_diagram",
                "visible_labels": [],
                "components": [],
                "relationships": [],
                "corrected_block_type": None,
                "corrected_level": None,
                "claims": [],
                "unresolved": ["No material readable claims."],
            })

    class Interactions:
        def create(self, **kwargs):
            requests.append(kwargs)
            prompt = kwargs["input"][0]["text"]
            candidate_id = "ai-0001" if "ai-0001" in prompt else "ai-0002"
            return Response(candidate_id)

    class Client:
        interactions = Interactions()

    monkeypatch.setattr(genai, "Client", lambda **kwargs: Client())
    pictures = tmp_path / "pictures"
    pictures.mkdir()
    candidates = []
    for index in range(1, 4):
        name = f"picture-{index:04d}.png"
        (pictures / name).write_bytes(f"image-{index}".encode())
        candidates.append({
            "id": f"ai-{index:04d}",
            "kind": "diagram",
            "block_id": f"block-{index}",
            "source_refs": [f"#/pictures/{index - 1}"],
            "section_path": ["Architecture"],
            "picture_file": name,
        })
    events = []
    settings = GeminiSettings(
        True,
        "secret",
        "gemini-3.5-flash",
        30,
        0,
        1,
        thinking_level="medium",
        budget_inr=100,
        max_candidates=2,
        max_requests=7,
        max_verifications=2,
        request_reserve_inr=0,
    )

    report, results, supplements = run_gemini_review(
        candidates,
        tmp_path,
        settings,
        tmp_path / "cache",
        events.append,
    )

    assert len(requests) == 2
    assert all(
        request["generation_config"]
        == {"thinking_level": "medium", "max_output_tokens": 8_192}
        for request in requests
    )
    assert report["selected_candidate_count"] == 2
    assert report["deferred_candidate_count"] == 1
    assert report["request_count"] == 2
    assert report["extraction_request_count"] == 2
    assert report["verification_request_count"] == 0
    assert report["diagram_request_count"] == 2
    assert report["submitted_evidence_count"] == 2
    assert report["verification_count"] == 0
    assert results["ai-0001"]["verification_skipped"] == "no_claims"
    assert results["ai-0003"]["status"] == "not_selected"
    assert candidates[2].get("submitted_evidence_paths") is None
    assert len(list((tmp_path / "ai_evidence").iterdir())) == 2
    assert any(event["event"] == "request_completed" for event in events)
    assert (tmp_path / "gemini_events.jsonl").is_file()
    assert supplements == []


def test_malformed_paid_response_records_usage_cost_raw_output_and_failure_events(
    monkeypatch,
    tmp_path: Path,
) -> None:
    from google import genai

    class Usage:
        def model_dump(self, **kwargs):
            return {
                "total_input_tokens": 1_000,
                "total_output_tokens": 2_000,
                "total_thought_tokens": 500,
                "total_tokens": 3_500,
            }

    class Response:
        usage = Usage()
        status = "in_progress"
        incomplete_details = None
        output_text = '{"candidate_id":"ai-0001","claims":[{"statement":"The'

    class Interactions:
        def create(self, **kwargs):
            return Response()

    class Client:
        interactions = Interactions()

    monkeypatch.setattr(genai, "Client", lambda **kwargs: Client())
    pictures = tmp_path / "pictures"
    pictures.mkdir()
    (pictures / "picture-0001.png").write_bytes(b"not-a-real-image")
    candidate = {
        "id": "ai-0001",
        "kind": "diagram",
        "block_id": "block-1",
        "source_refs": ["#/pictures/0"],
        "section_path": ["Architecture"],
        "picture_file": "picture-0001.png",
    }
    events = []
    settings = GeminiSettings(
        True,
        "secret",
        "gemini-3.5-flash",
        30,
        0,
        1,
        budget_inr=100,
        max_candidates=1,
        request_reserve_inr=0,
    )

    report, results, supplements = run_gemini_review(
        [candidate],
        tmp_path,
        settings,
        tmp_path / "cache",
        events.append,
    )

    result = results["ai-0001"]
    assert report["status"] == "failed"
    assert report["usage"]["total_tokens"] == 3_500
    assert report["estimated_cost_inr"] > 0
    assert report["submitted_candidate_count"] == 1
    assert report["submitted_evidence_count"] == 1
    assert result["status"] == "failed"
    assert result["usage"]["extraction"]["total_tokens"] == 3_500
    assert result["raw_extraction"].endswith('"The')
    assert any(event["event"] == "response_received" for event in events)
    assert any(event["event"] == "response_validation_failed" for event in events)
    assert not any(event["event"] == "request_completed" for event in events)
    assert supplements == []


def test_docx_picture_uses_parent_heading_context_in_document_order(tmp_path: Path) -> None:
    source = tmp_path / "anchored.docx"
    document = Document()
    document.add_heading("5.3. Frontend Architecture", level=2)
    document.add_paragraph("Architecture details follow.")
    document.save(source)
    document_data = {
        "texts": [
            {
                "self_ref": "#/texts/0",
                "label": "section_header",
                "text": "5.3. Frontend Architecture",
            },
            {
                "self_ref": "#/texts/1",
                "label": "text",
                "text": "Architecture details follow.",
            },
        ],
        "tables": [],
        "pictures": [{
            "self_ref": "#/pictures/0",
            "parent": {"$ref": "#/texts/0"},
            "image": {"uri": "assets/frontend.png"},
            "prov": [],
        }],
    }

    result = build_canonical_document(source, document_data, {}, {}, [])
    blocks = result.document["blocks"]
    picture = next(block for block in blocks if block["type"] == "picture")
    heading_index = next(i for i, block in enumerate(blocks) if block["type"] == "heading")
    picture_index = blocks.index(picture)

    assert picture_index == heading_index + 1
    assert picture["section_path"] == ["5.3. Frontend Architecture"]
    assert picture["native"]["anchor_ref"] == "#/texts/0"
    assert picture["native"]["anchor_matched"] is True


def test_docx_native_drawing_order_overrides_stale_docling_parent(tmp_path: Path) -> None:
    from PIL import Image

    source = tmp_path / "native-anchor.docx"
    image_path = tmp_path / "pipeline.png"
    Image.new("RGB", (120, 40), "white").save(image_path)
    document = Document()
    document.add_heading("10. Deployment", level=1)
    numeric_heading = document.add_paragraph()
    numeric_heading.add_run("10.4. CI/CD Pipeline").bold = True
    document.add_picture(str(image_path))
    document.save(source)
    document_data = {
        "texts": [
            {"self_ref": "#/texts/0", "text": "10. Deployment"},
            {"self_ref": "#/texts/1", "text": "10.4. CI/CD Pipeline"},
            {"self_ref": "#/texts/99", "text": "Stale unrelated parent"},
        ],
        "tables": [],
        "pictures": [{
            "self_ref": "#/pictures/0",
            "parent": {"$ref": "#/texts/99"},
            "image": {"uri": "assets/pipeline.png"},
            "prov": [],
        }],
    }

    result = build_canonical_document(source, document_data, {}, {}, [])
    blocks = result.document["blocks"]
    heading = next(block for block in blocks if "CI/CD Pipeline" in block.get("text", ""))
    picture = next(block for block in blocks if block["type"] == "picture")

    assert heading["type"] == "heading"
    assert heading["level"] == 2
    assert picture["section_path"] == ["10. Deployment", "10.4. CI/CD Pipeline"]
    assert picture["native"]["source"] == "docx_ooxml_drawing_order"


def test_provider_incomplete_response_is_audited_after_usage_is_recorded(
    monkeypatch,
    tmp_path: Path,
) -> None:
    from google import genai

    class Usage:
        def model_dump(self, **kwargs):
            return {"total_output_tokens": 2_048, "total_tokens": 2_048}

    class Details:
        def model_dump(self, **kwargs):
            return {"reason": "max_output_tokens"}

    class Response:
        usage = Usage()
        status = "incomplete"
        incomplete_details = Details()
        output_text = '{"candidate_id":"ai-0001"'

    class Interactions:
        def create(self, **kwargs):
            return Response()

    class Client:
        interactions = Interactions()

    monkeypatch.setattr(genai, "Client", lambda **kwargs: Client())
    candidate = {
        "id": "ai-0001",
        "kind": "structural_ambiguity",
        "block_id": "block-1",
        "source_refs": ["#/texts/0"],
        "section_path": ["Architecture"],
    }
    events = []
    settings = GeminiSettings(True, "secret", "gemini-3.5-flash", 30, 0, 1)

    report, results, _ = run_gemini_review(
        [candidate], tmp_path, settings, tmp_path / "cache", events.append
    )

    assert report["status"] == "failed"
    assert report["usage"]["total_output_tokens"] == 2_048
    assert results["ai-0001"]["response_status"] == "incomplete"
    assert results["ai-0001"]["incomplete_details"] == {"reason": "max_output_tokens"}
    assert any(event["event"] == "response_incomplete" for event in events)


def test_complex_erd_submission_keeps_full_image_and_adds_targeted_tiles(
    monkeypatch,
    tmp_path: Path,
) -> None:
    from google import genai
    from PIL import Image

    class Response:
        usage = None
        output_text = json.dumps({
            "candidate_id": "ai-0001",
            "classification": "database_erd",
            "visible_labels": [],
            "components": [],
            "relationships": [],
            "corrected_block_type": None,
            "corrected_level": None,
            "claims": [],
            "unresolved": ["No bounded claim selected."],
        })

    class Interactions:
        def create(self, **kwargs):
            return Response()

    class Client:
        interactions = Interactions()

    monkeypatch.setattr(genai, "Client", lambda **kwargs: Client())
    pictures = tmp_path / "pictures"
    pictures.mkdir()
    Image.new("RGB", (1_600, 1_000), "white").save(pictures / "picture-0001.png")
    candidate = {
        "id": "ai-0001",
        "kind": "diagram",
        "block_id": "block-1",
        "source_refs": ["#/pictures/0"],
        "section_path": ["Supplier ER Diagram"],
        "picture_file": "picture-0001.png",
        "ocr_hint": "uuid supplier_id FK enum varchar timestamp",
    }

    report, results, _ = run_gemini_review(
        [candidate],
        tmp_path,
        GeminiSettings(True, "secret", "gemini-3.5-flash", 30, 0, 1),
        tmp_path / "cache",
    )

    assert candidate["evidence_strategy"] == "full_plus_erd_tiles"
    assert len(candidate["prepared_evidence_paths"]) == 3
    assert len(candidate["submitted_evidence_paths"]) == 3
    assert report["submitted_evidence_count"] == 3
    assert results["ai-0001"]["status"] == "unresolved"


def test_ai_review_defaults_off_when_environment_is_absent(monkeypatch) -> None:
    monkeypatch.delenv("DECIDIAN_AI_REVIEW", raising=False)

    settings = get_gemini_settings()

    assert settings.enabled is False


def test_rupee_budget_stops_before_verification(monkeypatch, tmp_path: Path) -> None:
    from google import genai

    class Usage:
        def model_dump(self, **kwargs):
            return {
                "total_input_tokens": 10_000,
                "total_output_tokens": 5_000,
                "total_thought_tokens": 5_000,
                "total_tokens": 20_000,
            }

    class Response:
        usage = Usage()
        output_text = json.dumps({
            "candidate_id": "ai-0001",
            "classification": "architecture_diagram",
            "visible_labels": [],
            "components": ["API"],
            "relationships": [],
            "corrected_block_type": None,
            "corrected_level": None,
            "claims": [{
                "claim_id": "c-0001",
                "claim_type": "component",
                "statement": "An API component is shown.",
                "evidence": [{"label": "API", "box_2d": [1, 1, 10, 10]}],
            }],
            "unresolved": [],
        })

    calls = []

    class Interactions:
        def create(self, **kwargs):
            calls.append(kwargs)
            return Response()

    class Client:
        interactions = Interactions()

    monkeypatch.setattr(genai, "Client", lambda **kwargs: Client())
    settings = GeminiSettings(
        True,
        "secret",
        "gemini-3.5-flash",
        30,
        0,
        1,
        budget_inr=5,
        usd_inr_rate=100,
        input_price_usd_per_million=1.5,
        output_price_usd_per_million=9,
        max_candidates=1,
        max_requests=7,
        request_reserve_inr=0,
    )
    candidate = {
        "id": "ai-0001",
        "kind": "diagram",
        "block_id": "block-1",
        "source_refs": ["#/pictures/0"],
        "section_path": ["Architecture"],
    }

    report, results, supplements = run_gemini_review(
        [candidate],
        tmp_path,
        settings,
        tmp_path / "cache",
    )

    assert len(calls) == 1
    assert report["request_count"] == 1
    assert report["stop_reason"] == "budget_limit"
    assert report["estimated_cost_inr"] > settings.budget_inr
    assert results["ai-0001"]["status"] == "guard_stopped"
    assert results["ai-0001"]["stop_reason"] == "budget_limit"
    assert supplements == []
